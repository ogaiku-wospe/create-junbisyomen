"""
全形式対応ファイルプロセッサー
- JPEG/PNG/HEIC: 画像処理
- PDF: テキスト抽出・OCR
- Word/Excel: 文書処理
- HTML: ウェブページ解析
- MP4/MP3: 動画・音声処理
"""

import os
import json
import logging
import subprocess
from typing import Dict, List, Optional, Tuple
from pathlib import Path

# 画像処理
try:
    from PIL import Image
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

# PDF処理
try:
    import PyPDF2
    from pdf2image import convert_from_path
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word処理
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Excel処理
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# HTML処理
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# OCR
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

from global_config import *

logger = logging.getLogger(__name__)


class FileProcessor:
    """全形式対応ファイルプロセッサー"""
    
    def __init__(self):
        """初期化"""
        self.temp_dir = LOCAL_TEMP_DIR
        os.makedirs(self.temp_dir, exist_ok=True)
        logger.info("✅ FileProcessor初期化完了")
    
    def process_file(self, file_path: str, file_type: str) -> Dict:
        """
        ファイルを処理して内容を抽出
        
        Args:
            file_path: ファイルパス
            file_type: ファイルタイプ
        
        Returns:
            抽出結果
        """
        logger.info(f"🔄 ファイル処理開始: {os.path.basename(file_path)} (タイプ: {file_type})")
        
        # ファイルタイプに応じた処理
        if file_type == 'image':
            return self._process_image(file_path)
        elif file_type == 'pdf':
            return self._process_pdf(file_path)
        elif file_type == 'document':
            return self._process_document(file_path)
        elif file_type == 'spreadsheet':
            return self._process_spreadsheet(file_path)
        elif file_type == 'presentation':
            return self._process_presentation(file_path)
        elif file_type == 'web':
            return self._process_html(file_path)
        elif file_type == 'video':
            return self._process_video(file_path)
        elif file_type == 'audio':
            return self._process_audio(file_path)
        elif file_type == 'email':
            return self._process_email(file_path)
        elif file_type == 'archive':
            return self._process_archive(file_path)
        else:
            return self._process_generic(file_path)
    
    # ================================
    # 画像処理
    # ================================
    
    def _process_image(self, image_path: str) -> Dict:
        """
        画像を処理
        - JPEG, PNG, HEIC等に対応
        - OCR実行
        - 画像情報抽出
        """
        logger.info(f"🖼️ 画像処理: {os.path.basename(image_path)}")
        
        result = {
            "file_type": "image",
            "processing_status": "success",
            "content": {}
        }
        
        try:
            # HEIC変換
            if image_path.lower().endswith(('.heic', '.heif')):
                image_path = self._convert_heic_to_jpg(image_path)
            
            # 画像読み込み
            if PILLOW_AVAILABLE:
                with Image.open(image_path) as img:
                    result['content']['image_info'] = {
                        "format": img.format,
                        "size": img.size,
                        "mode": img.mode
                    }
                    
                    # OCR実行
                    if OCR_ENABLED and TESSERACT_AVAILABLE:
                        ocr_result = self._perform_ocr(image_path)
                        result['content']['ocr_text'] = ocr_result['text']
                        result['content']['ocr_confidence'] = ocr_result['confidence']
                        result['content']['ocr_language'] = ocr_result['language']
            
            # 画像を一時保存（AI分析用）
            result['processed_file_path'] = image_path
            
            logger.info(f"✅ 画像処理完了")
            return result
            
        except Exception as e:
            logger.error(f"❌ 画像処理失敗: {e}")
            result['processing_status'] = 'error'
            result['error'] = str(e)
            return result
    
    def _convert_heic_to_jpg(self, heic_path: str) -> str:
        """HEICをJPGに変換"""
        try:
            output_path = os.path.join(
                self.temp_dir,
                f"{Path(heic_path).stem}_converted.jpg"
            )
            
            # PIL での変換
            if PILLOW_AVAILABLE:
                with Image.open(heic_path) as img:
                    rgb_img = img.convert('RGB')
                    rgb_img.save(output_path, 'JPEG', quality=95)
                
                logger.info(f"✅ HEIC変換完了: {output_path}")
                return output_path
            else:
                logger.warning("⚠️ PIL未インストール - HEIC変換スキップ")
                return heic_path
                
        except Exception as e:
            logger.error(f"❌ HEIC変換失敗: {e}")
            return heic_path
    
    def _perform_ocr(self, image_path: str) -> Dict:
        """OCR実行"""
        try:
            # Tesseract OCR
            text = pytesseract.image_to_string(
                Image.open(image_path),
                lang='+'.join(OCR_LANGUAGES)
            )
            
            # 詳細情報取得
            data = pytesseract.image_to_data(
                Image.open(image_path),
                lang='+'.join(OCR_LANGUAGES),
                output_type=pytesseract.Output.DICT
            )
            
            # 信頼度計算
            confidences = [int(conf) for conf in data['conf'] if conf != '-1']
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            return {
                "text": text.strip(),
                "confidence": round(avg_confidence / 100, 2),
                "language": OCR_LANGUAGES[0],
                "word_count": len(text.split()),
                "char_count": len(text)
            }
            
        except Exception as e:
            logger.error(f"❌ OCR実行失敗: {e}")
            return {
                "text": "",
                "confidence": 0.0,
                "language": "unknown",
                "error": str(e)
            }
    
    # ================================
    # PDF処理
    # ================================
    
    def _process_pdf(self, pdf_path: str) -> Dict:
        """
        PDF処理
        - テキスト抽出
        - 画像変換してOCR
        """
        logger.info(f"📄 PDF処理: {os.path.basename(pdf_path)}")
        
        result = {
            "file_type": "pdf",
            "processing_status": "success",
            "content": {}
        }
        
        try:
            if not PDF_AVAILABLE:
                raise ImportError("PyPDF2未インストール")
            
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # 基本情報
                result['content']['page_count'] = len(pdf_reader.pages)
                result['content']['is_encrypted'] = pdf_reader.is_encrypted
                
                # 全ページのテキスト抽出
                full_text = []
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    full_text.append({
                        "page_number": i + 1,
                        "text": page_text,
                        "char_count": len(page_text)
                    })
                
                result['content']['pages'] = full_text
                result['content']['total_text'] = '\n\n'.join([p['text'] for p in full_text])
                
                # テキストが少ない場合はOCR実行
                total_chars = sum(p['char_count'] for p in full_text)
                if total_chars < 100 and OCR_ENABLED:
                    logger.info("📷 PDF→画像変換してOCR実行")
                    ocr_result = self._pdf_to_image_ocr(pdf_path)
                    result['content']['ocr_results'] = ocr_result
            
            result['processed_file_path'] = pdf_path
            logger.info(f"✅ PDF処理完了")
            return result
            
        except Exception as e:
            logger.error(f"❌ PDF処理失敗: {e}")
            result['processing_status'] = 'error'
            result['error'] = str(e)
            return result
    
    def _pdf_to_image_ocr(self, pdf_path: str) -> List[Dict]:
        """PDFを画像化してOCR"""
        try:
            # PDF→画像変換
            images = convert_from_path(pdf_path, dpi=300)
            
            ocr_results = []
            for i, image in enumerate(images):
                # 一時保存
                temp_image_path = os.path.join(
                    self.temp_dir,
                    f"pdf_page_{i+1}.jpg"
                )
                image.save(temp_image_path, 'JPEG')
                
                # OCR実行
                ocr_result = self._perform_ocr(temp_image_path)
                ocr_results.append({
                    "page_number": i + 1,
                    "ocr_text": ocr_result['text'],
                    "confidence": ocr_result['confidence']
                })
                
                # 一時ファイル削除
                os.remove(temp_image_path)
            
            return ocr_results
            
        except Exception as e:
            logger.error(f"❌ PDF→画像OCR失敗: {e}")
            return []
    
    # ================================
    # Word文書処理
    # ================================
    
    def _process_document(self, doc_path: str) -> Dict:
        """Word文書処理"""
        logger.info(f"📝 Word文書処理: {os.path.basename(doc_path)}")
        
        result = {
            "file_type": "document",
            "processing_status": "success",
            "content": {}
        }
        
        try:
            if doc_path.lower().endswith('.docx'):
                return self._process_docx(doc_path)
            else:
                # .docの場合はPDFに変換してから処理
                logger.warning("⚠️ .doc形式は.docxに変換が必要")
                result['processing_status'] = 'partial'
                result['error'] = '.doc形式の完全サポートには変換が必要'
                return result
                
        except Exception as e:
            logger.error(f"❌ Word文書処理失敗: {e}")
            result['processing_status'] = 'error'
            result['error'] = str(e)
            return result
    
    def _process_docx(self, docx_path: str) -> Dict:
        """DOCX処理"""
        if not DOCX_AVAILABLE:
            return {
                "file_type": "document",
                "processing_status": "error",
                "error": "python-docx未インストール"
            }
        
        try:
            doc = Document(docx_path)
            
            # 段落抽出
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        "paragraph_number": i + 1,
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })
            
            # 表抽出
            tables = []
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    "table_number": i + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": table_data
                })
            
            # 全文結合
            full_text = '\n\n'.join([p['text'] for p in paragraphs])
            
            return {
                "file_type": "document",
                "processing_status": "success",
                "content": {
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "full_text": full_text,
                    "paragraph_count": len(paragraphs),
                    "table_count": len(tables),
                    "word_count": len(full_text.split()),
                    "char_count": len(full_text)
                },
                "processed_file_path": docx_path
            }
            
        except Exception as e:
            logger.error(f"❌ DOCX処理失敗: {e}")
            return {
                "file_type": "document",
                "processing_status": "error",
                "error": str(e)
            }
    
    # ================================
    # Excel処理
    # ================================
    
    def _process_spreadsheet(self, xlsx_path: str) -> Dict:
        """Excel処理"""
        logger.info(f"📊 Excel処理: {os.path.basename(xlsx_path)}")
        
        result = {
            "file_type": "spreadsheet",
            "processing_status": "success",
            "content": {}
        }
        
        try:
            if not EXCEL_AVAILABLE:
                raise ImportError("openpyxl未インストール")
            
            workbook = openpyxl.load_workbook(xlsx_path, data_only=True)
            
            sheets_data = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # シートデータ抽出
                rows_data = []
                for row in sheet.iter_rows(values_only=True):
                    rows_data.append(list(row))
                
                sheets_data.append({
                    "sheet_name": sheet_name,
                    "rows": len(rows_data),
                    "columns": len(rows_data[0]) if rows_data else 0,
                    "data": rows_data[:100]  # 最大100行
                })
            
            result['content']['sheets'] = sheets_data
            result['content']['sheet_count'] = len(sheets_data)
            result['processed_file_path'] = xlsx_path
            
            logger.info(f"✅ Excel処理完了: {len(sheets_data)}シート")
            return result
            
        except Exception as e:
            logger.error(f"❌ Excel処理失敗: {e}")
            result['processing_status'] = 'error'
            result['error'] = str(e)
            return result
    
    # ================================
    # プレゼンテーション処理
    # ================================
    
    def _process_presentation(self, pptx_path: str) -> Dict:
        """PowerPoint処理"""
        logger.info(f"📊 PowerPoint処理: {os.path.basename(pptx_path)}")
        
        # TODO: python-pptx使用
        return {
            "file_type": "presentation",
            "processing_status": "partial",
            "error": "PowerPoint処理は今後実装"
        }
    
    # ================================
    # HTML処理
    # ================================
    
    def _process_html(self, html_path: str) -> Dict:
        """HTML処理"""
        logger.info(f"🌐 HTML処理: {os.path.basename(html_path)}")
        
        result = {
            "file_type": "web",
            "processing_status": "success",
            "content": {}
        }
        
        try:
            if not BS4_AVAILABLE:
                raise ImportError("BeautifulSoup未インストール")
            
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # テキスト抽出
            text = soup.get_text(separator='\n', strip=True)
            
            # メタ情報抽出
            meta_tags = {}
            for meta in soup.find_all('meta'):
                name = meta.get('name') or meta.get('property')
                content = meta.get('content')
                if name and content:
                    meta_tags[name] = content
            
            # リンク抽出
            links = []
            for link in soup.find_all('a', href=True):
                links.append({
                    "text": link.get_text(strip=True),
                    "href": link['href']
                })
            
            result['content'] = {
                "title": soup.title.string if soup.title else None,
                "meta_tags": meta_tags,
                "text": text,
                "links": links[:100],  # 最大100リンク
                "link_count": len(links)
            }
            
            result['processed_file_path'] = html_path
            logger.info(f"✅ HTML処理完了")
            return result
            
        except Exception as e:
            logger.error(f"❌ HTML処理失敗: {e}")
            result['processing_status'] = 'error'
            result['error'] = str(e)
            return result
    
    # ================================
    # 動画処理
    # ================================
    
    def _process_video(self, video_path: str) -> Dict:
        """動画処理"""
        logger.info(f"🎬 動画処理: {os.path.basename(video_path)}")
        
        result = {
            "file_type": "video",
            "processing_status": "success",
            "content": {}
        }
        
        try:
            # ffprobeで情報取得
            video_info = self._get_video_info(video_path)
            result['content']['video_info'] = video_info
            
            # フレーム抽出
            frames = self._extract_video_frames(video_path)
            result['content']['extracted_frames'] = frames
            
            # 音声抽出
            audio_path = self._extract_audio_from_video(video_path)
            if audio_path:
                # 音声の文字起こし
                transcription = self._transcribe_audio(audio_path)
                result['content']['transcription'] = transcription
            
            result['processed_file_path'] = video_path
            logger.info(f"✅ 動画処理完了")
            return result
            
        except Exception as e:
            logger.error(f"❌ 動画処理失敗: {e}")
            result['processing_status'] = 'error'
            result['error'] = str(e)
            return result
    
    def _get_video_info(self, video_path: str) -> Dict:
        """動画情報取得"""
        try:
            cmd = [
                'ffprobe',
                '-v', 'quiet',
                '-print_format', 'json',
                '-show_format',
                '-show_streams',
                video_path
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                return json.loads(result.stdout)
            else:
                return {"error": "ffprobe実行失敗"}
                
        except Exception as e:
            logger.error(f"❌ 動画情報取得失敗: {e}")
            return {"error": str(e)}
    
    def _extract_video_frames(self, video_path: str) -> List[str]:
        """動画からフレーム抽出"""
        # TODO: OpenCV使用
        logger.warning("⚠️ フレーム抽出は今後実装")
        return []
    
    def _extract_audio_from_video(self, video_path: str) -> Optional[str]:
        """動画から音声抽出"""
        try:
            output_path = os.path.join(
                self.temp_dir,
                f"{Path(video_path).stem}_audio.mp3"
            )
            
            cmd = [
                'ffmpeg',
                '-i', video_path,
                '-vn',
                '-acodec', 'libmp3lame',
                '-y',
                output_path
            ]
            
            result = subprocess.run(cmd, capture_output=True)
            if result.returncode == 0:
                return output_path
            else:
                return None
                
        except Exception as e:
            logger.error(f"❌ 音声抽出失敗: {e}")
            return None
    
    # ================================
    # 音声処理
    # ================================
    
    def _process_audio(self, audio_path: str) -> Dict:
        """音声処理"""
        logger.info(f"🎵 音声処理: {os.path.basename(audio_path)}")
        
        result = {
            "file_type": "audio",
            "processing_status": "success",
            "content": {}
        }
        
        try:
            # 音声情報取得
            audio_info = self._get_audio_info(audio_path)
            result['content']['audio_info'] = audio_info
            
            # 文字起こし
            transcription = self._transcribe_audio(audio_path)
            result['content']['transcription'] = transcription
            
            result['processed_file_path'] = audio_path
            logger.info(f"✅ 音声処理完了")
            return result
            
        except Exception as e:
            logger.error(f"❌ 音声処理失敗: {e}")
            result['processing_status'] = 'error'
            result['error'] = str(e)
            return result
    
    def _get_audio_info(self, audio_path: str) -> Dict:
        """音声情報取得"""
        # TODO: ffprobe使用
        return {"status": "not_implemented"}
    
    def _transcribe_audio(self, audio_path: str) -> Dict:
        """音声文字起こし"""
        # TODO: Whisper使用
        logger.warning("⚠️ 音声文字起こしは今後実装（Whisper使用）")
        return {
            "text": "",
            "language": "unknown",
            "confidence": 0.0,
            "status": "not_implemented"
        }
    
    # ================================
    # メール処理
    # ================================
    
    def _process_email(self, email_path: str) -> Dict:
        """メール処理"""
        logger.info(f"📧 メール処理: {os.path.basename(email_path)}")
        
        # TODO: email.parser使用
        return {
            "file_type": "email",
            "processing_status": "partial",
            "error": "メール処理は今後実装"
        }
    
    # ================================
    # アーカイブ処理
    # ================================
    
    def _process_archive(self, archive_path: str) -> Dict:
        """アーカイブ処理"""
        logger.info(f"📦 アーカイブ処理: {os.path.basename(archive_path)}")
        
        # TODO: zipfile/rarfile使用
        return {
            "file_type": "archive",
            "processing_status": "partial",
            "error": "アーカイブ処理は今後実装"
        }
    
    # ================================
    # 汎用処理
    # ================================
    
    def _process_generic(self, file_path: str) -> Dict:
        """汎用ファイル処理"""
        logger.warning(f"⚠️ 未対応形式: {os.path.basename(file_path)}")
        
        return {
            "file_type": "unknown",
            "processing_status": "unsupported",
            "error": "未対応のファイル形式",
            "content": {}
        }
