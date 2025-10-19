"""
完全なメタデータ抽出モジュール
- EXIF情報
- ファイルハッシュ
- Google Drive URL
- バージョン情報
"""

import os
import hashlib
import mimetypes
import json
from typing import Dict, Optional, Any
from datetime import datetime
from pathlib import Path

try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from config import *
import logging

logger = logging.getLogger(__name__)


class MetadataExtractor:
    """完全なメタデータ抽出クラス"""
    
    def __init__(self):
        """初期化"""
        self.hash_algorithm = HASH_ALGORITHM
        logger.info("✅ MetadataExtractor初期化完了")
    
    def extract_complete_metadata(self, 
                                  file_path: str, 
                                  gdrive_file_info: Dict = None) -> Dict:
        """
        完全なメタデータを抽出
        
        Args:
            file_path: ファイルパス
            gdrive_file_info: Google Driveのファイル情報
        
        Returns:
            完全なメタデータ辞書
        """
        logger.info(f"📊 メタデータ抽出開始: {os.path.basename(file_path)}")
        
        metadata = {
            # 基本情報
            "basic": self._extract_basic_metadata(file_path),
            
            # ファイルハッシュ
            "hashes": self._calculate_hashes(file_path),
            
            # Google Drive情報
            "gdrive": self._extract_gdrive_metadata(gdrive_file_info),
            
            # ファイル形式固有情報
            "format_specific": self._extract_format_specific_metadata(file_path),
            
            # 抽出メタ情報
            "extraction_info": {
                "timestamp": get_timestamp(),
                "extractor_version": SYSTEM_VERSION,
                "extraction_method": "complete"
            }
        }
        
        logger.info(f"✅ メタデータ抽出完了")
        return metadata
    
    def _extract_basic_metadata(self, file_path: str) -> Dict:
        """基本メタデータを抽出"""
        try:
            stat = os.stat(file_path)
            file_name = os.path.basename(file_path)
            file_ext = os.path.splitext(file_name)[1].lower()
            
            # MIMEタイプ判定
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                mime_type = self._detect_mime_type(file_path)
            
            return {
                "file_name": file_name,
                "file_extension": file_ext,
                "file_size_bytes": stat.st_size,
                "file_size_human": self._format_file_size(stat.st_size),
                "mime_type": mime_type,
                "created_time": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed_time": datetime.fromtimestamp(stat.st_atime).isoformat(),
                "file_permissions": oct(stat.st_mode)[-3:],
                "is_readable": os.access(file_path, os.R_OK),
                "is_writable": os.access(file_path, os.W_OK)
            }
        except Exception as e:
            logger.error(f"❌ 基本メタデータ抽出失敗: {e}")
            return {}
    
    def _calculate_hashes(self, file_path: str) -> Dict:
        """複数のハッシュ値を計算"""
        try:
            hashes = {}
            
            # SHA-256
            sha256 = hashlib.sha256()
            # MD5
            md5 = hashlib.md5()
            # SHA-1
            sha1 = hashlib.sha1()
            
            with open(file_path, 'rb') as f:
                while chunk := f.read(8192):
                    sha256.update(chunk)
                    md5.update(chunk)
                    sha1.update(chunk)
            
            hashes = {
                "sha256": sha256.hexdigest(),
                "md5": md5.hexdigest(),
                "sha1": sha1.hexdigest(),
                "algorithm_primary": "sha256"
            }
            
            logger.debug(f"ハッシュ計算完了: SHA-256={hashes['sha256'][:16]}...")
            return hashes
            
        except Exception as e:
            logger.error(f"❌ ハッシュ計算失敗: {e}")
            return {}
    
    def _extract_gdrive_metadata(self, gdrive_file_info: Dict = None) -> Dict:
        """Google Driveメタデータを抽出"""
        if not gdrive_file_info:
            return {}
        
        try:
            file_id = gdrive_file_info.get('id', '')
            
            metadata = {
                "file_id": file_id,
                "file_url": GDRIVE_FILE_URL_FORMAT.format(file_id=file_id),
                "download_url": f"https://drive.google.com/uc?export=download&id={file_id}",
                "web_view_link": gdrive_file_info.get('webViewLink', ''),
                "web_content_link": gdrive_file_info.get('webContentLink', ''),
                "thumbnail_link": gdrive_file_info.get('thumbnailLink', ''),
                "icon_link": gdrive_file_info.get('iconLink', ''),
                "parent_folders": gdrive_file_info.get('parents', []),
                "owners": gdrive_file_info.get('owners', []),
                "last_modifying_user": gdrive_file_info.get('lastModifyingUser', {}),
                "created_time": gdrive_file_info.get('createdTime', ''),
                "modified_time": gdrive_file_info.get('modifiedTime', ''),
                "version": gdrive_file_info.get('version', ''),
                "original_filename": gdrive_file_info.get('originalFilename', ''),
                "md5_checksum": gdrive_file_info.get('md5Checksum', ''),
                "shared": gdrive_file_info.get('shared', False),
                "capabilities": gdrive_file_info.get('capabilities', {})
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"❌ Google Driveメタデータ抽出失敗: {e}")
            return {}
    
    def _extract_format_specific_metadata(self, file_path: str) -> Dict:
        """ファイル形式固有のメタデータを抽出"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # 画像ファイル
        if file_ext in SUPPORTED_FORMATS['image']['extensions']:
            return self._extract_image_metadata(file_path)
        
        # PDFファイル
        elif file_ext == '.pdf':
            return self._extract_pdf_metadata(file_path)
        
        # Word文書
        elif file_ext in ['.docx', '.doc']:
            return self._extract_docx_metadata(file_path)
        
        # その他
        else:
            return {"format": "unsupported_format_specific"}
    
    def _extract_image_metadata(self, image_path: str) -> Dict:
        """画像のEXIF情報を抽出"""
        if not PILLOW_AVAILABLE:
            return {"error": "PIL not available"}
        
        try:
            with Image.open(image_path) as img:
                metadata = {
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "width": img.width,
                    "height": img.height,
                    "aspect_ratio": round(img.width / img.height, 2) if img.height > 0 else 0,
                    "resolution_dpi": img.info.get('dpi', None),
                    "color_space": self._get_color_space(img.mode),
                    "has_transparency": img.mode in ('RGBA', 'LA', 'P')
                }
                
                # EXIF情報
                exif_data = img._getexif() if hasattr(img, '_getexif') else None
                if exif_data:
                    exif = {}
                    for tag_id, value in exif_data.items():
                        tag = TAGS.get(tag_id, tag_id)
                        
                        # GPS情報の特別処理
                        if tag == "GPSInfo":
                            gps_data = {}
                            for t in value:
                                sub_tag = GPSTAGS.get(t, t)
                                gps_data[sub_tag] = value[t]
                            exif[tag] = gps_data
                        else:
                            # バイナリデータは文字列化
                            if isinstance(value, bytes):
                                try:
                                    value = value.decode('utf-8')
                                except:
                                    value = str(value)
                            exif[tag] = str(value)
                    
                    metadata['exif'] = exif
                
                return metadata
                
        except Exception as e:
            logger.error(f"❌ 画像メタデータ抽出失敗: {e}")
            return {"error": str(e)}
    
    def _extract_pdf_metadata(self, pdf_path: str) -> Dict:
        """PDFのメタデータを抽出"""
        if not PYPDF2_AVAILABLE:
            return {"error": "PyPDF2 not available"}
        
        try:
            with open(pdf_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "is_encrypted": pdf_reader.is_encrypted,
                    "pdf_version": pdf_reader.pdf_header if hasattr(pdf_reader, 'pdf_header') else None
                }
                
                # ドキュメント情報
                if pdf_reader.metadata:
                    info = {}
                    for key, value in pdf_reader.metadata.items():
                        # PyPDF2のキーは'/Title'等の形式
                        clean_key = key.lstrip('/')
                        info[clean_key] = str(value) if value else None
                    metadata['document_info'] = info
                
                # 最初のページのサイズ
                if len(pdf_reader.pages) > 0:
                    page = pdf_reader.pages[0]
                    metadata['page_size'] = {
                        'width': float(page.mediabox.width) if hasattr(page, 'mediabox') else None,
                        'height': float(page.mediabox.height) if hasattr(page, 'mediabox') else None
                    }
                
                return metadata
                
        except Exception as e:
            logger.error(f"❌ PDFメタデータ抽出失敗: {e}")
            return {"error": str(e)}
    
    def _extract_docx_metadata(self, docx_path: str) -> Dict:
        """Word文書のメタデータを抽出"""
        if not DOCX_AVAILABLE:
            return {"error": "python-docx not available"}
        
        try:
            doc = Document(docx_path)
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "table_count": len(doc.tables),
                "style_count": len(doc.styles)
            }
            
            # コアプロパティ
            core_props = doc.core_properties
            metadata['properties'] = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "comments": core_props.comments,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
                "category": core_props.category
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"❌ Word文書メタデータ抽出失敗: {e}")
            return {"error": str(e)}
    
    # ================================
    # ユーティリティメソッド
    # ================================
    
    def _detect_mime_type(self, file_path: str) -> str:
        """MIMEタイプを検出"""
        ext = os.path.splitext(file_path)[1].lower()
        
        # SUPPORTED_FORMATSから検索
        for format_type, format_info in SUPPORTED_FORMATS.items():
            if ext in format_info['extensions']:
                if format_info['mime_types']:
                    return format_info['mime_types'][0]
        
        return 'application/octet-stream'
    
    def _format_file_size(self, size_bytes: int) -> str:
        """ファイルサイズを人間が読める形式に変換"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.2f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} PB"
    
    def _get_color_space(self, mode: str) -> str:
        """PIL modeからカラースペースを取得"""
        color_spaces = {
            '1': 'Binary',
            'L': 'Grayscale',
            'P': 'Palette',
            'RGB': 'RGB',
            'RGBA': 'RGBA',
            'CMYK': 'CMYK',
            'YCbCr': 'YCbCr',
            'LAB': 'LAB',
            'HSV': 'HSV'
        }
        return color_spaces.get(mode, mode)
    
    def validate_metadata(self, metadata: Dict) -> Dict:
        """
        メタデータの完全性を検証
        
        Returns:
            検証結果
        """
        validation = {
            "is_valid": True,
            "completeness_score": 0.0,
            "missing_fields": [],
            "warnings": []
        }
        
        required_fields = [
            "basic.file_name",
            "basic.file_size_bytes",
            "basic.mime_type",
            "hashes.sha256",
            "gdrive.file_id",
            "gdrive.file_url"
        ]
        
        # フィールドの存在チェック
        present_count = 0
        for field in required_fields:
            parts = field.split('.')
            current = metadata
            exists = True
            
            for part in parts:
                if isinstance(current, dict) and part in current:
                    current = current[part]
                else:
                    exists = False
                    break
            
            if exists and current:
                present_count += 1
            else:
                validation['missing_fields'].append(field)
        
        validation['completeness_score'] = present_count / len(required_fields)
        
        if validation['completeness_score'] < QUALITY_CHECK_THRESHOLDS['metadata_coverage']:
            validation['is_valid'] = False
            validation['warnings'].append(
                f"メタデータカバレッジが不足: {validation['completeness_score']:.1%} < {QUALITY_CHECK_THRESHOLDS['metadata_coverage']:.1%}"
            )
        
        return validation
